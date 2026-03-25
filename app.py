"""
AI-Powered Resume Intelligence & Ranking System
"""

import streamlit as st
import pandas as pd
import io
import json
import re
from typing import List, Dict, Any
import sys
import os

# File processing imports (for resume upload)
PDF_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    pass

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    pass

# PDF generation imports
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    PDF_GENERATION_AVAILABLE = True
except ImportError:
    PDF_GENERATION_AVAILABLE = False

# Add src directory to path
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

# Force fallback mode to avoid torch issues entirely
USE_ORIGINAL = False


# Fallback standalone classes (only used if original modules fail)
class StandaloneRanker:
    """Fallback ranking system when original modules are unavailable."""
    
    def __init__(self):
        self.skills_db = {
            'Programming': ['python', 'java', 'c++', 'javascript', 'typescript', 'go', 'rust', 'c#', 'php', 'ruby'],
            'Machine Learning': ['tensorflow', 'pytorch', 'scikit-learn', 'keras', 'xgboost', 'pandas', 'numpy'],
            'Web Development': ['react', 'angular', 'vue', 'django', 'flask', 'nodejs', 'express'],
            'Databases': ['mysql', 'postgresql', 'mongodb', 'redis', 'sql', 'nosql'],
            'Cloud & DevOps': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git'],
            'Data Engineering': ['spark', 'hadoop', 'kafka', 'airflow', 'databricks'],
            'Mobile Development': ['react native', 'flutter', 'ios', 'android', 'swift'],
            'Project Management': ['agile', 'scrum', 'kanban', 'jira', 'confluence'],
            'Soft Skills': ['communication', 'leadership', 'teamwork', 'problem solving']
        }
        
        self.weights = {'skills': 0.6, 'experience': 0.25, 'education': 0.1, 'keywords': 0.05}
    
    def get_weights(self):
        """Get current scoring weights from session state or use defaults."""
        if hasattr(st, 'session_state') and 'scoring_weights' in st.session_state:
            weights = st.session_state.scoring_weights
            return {
                'skills': weights.get('skill_match', 60) / 100,
                'experience': weights.get('experience', 25) / 100,
                'education': 0.1,
                'keywords': weights.get('semantic_similarity', 15) / 100
            }
        return self.weights
    
    def analyze_job(self, job_text: str) -> Dict[str, Any]:
        skills = self._extract_skills(job_text)
        bias = self._detect_bias(job_text)
        experience = self._extract_experience(job_text)
        
        return {
            'text': job_text,
            'skills': skills,
            'bias_analysis': bias,
            'experience_requirements': {'minimum_years': experience}
        }
    
    def analyze_resume(self, resume_text: str, filename: str, job_analysis: Dict[str, Any], scoring_weights: Dict[str, Any] = None) -> Dict[str, Any]:
        contact = self._extract_contact(resume_text)
        skills = self._extract_skills(resume_text)
        experience = self._extract_experience(resume_text)
        bias = self._detect_bias(resume_text)
        
        skill_match = self._calculate_skill_match(job_analysis['skills'], skills)
        
        skill_score = skill_match['overall_coverage']
        experience_score = min(experience / max(job_analysis['experience_requirements']['minimum_years'], 1) * 100, 100)
        education_score = 75
        keyword_score = self._keyword_similarity(resume_text, job_analysis['text'])
        
        # Get current weights from configuration or use provided weights
        if scoring_weights:
            weights = {
                'skills': scoring_weights.get('skill_match', 60) / 100,
                'experience': scoring_weights.get('experience', 25) / 100,
                'education': 0.1,
                'keywords': scoring_weights.get('semantic_similarity', 15) / 100
            }
        else:
            weights = self.get_weights()
        
        final_score = (skill_score * weights['skills'] + 
                      experience_score * weights['experience'] + 
                      education_score * weights['education'] + 
                      keyword_score * weights['keywords'])
        
        return {
            'filename': filename,
            'name': contact['name'],
            'email': contact['email'],
            'phone': contact['phone'],
            'skill_coverage': skill_match,
            'years_of_experience': experience,
            'bias_analysis': bias,
            'final_score': final_score,
            'semantic_similarity': keyword_score  # Fallback
        }
    
    def rank_candidates(self, candidates: List[Dict[str, Any]], job_analysis: Dict[str, Any]) -> List[Dict[str, Any]]:
        candidates.sort(key=lambda x: x['final_score'], reverse=True)
        for i, candidate in enumerate(candidates, 1):
            candidate['rank'] = i
        return candidates
    
    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        text_lower = text.lower()
        found_skills = {}
        
        for category, skills in self.skills_db.items():
            category_skills = [skill for skill in skills if skill in text_lower]
            if category_skills:
                found_skills[category] = category_skills
        
        return found_skills
    
    def _extract_contact(self, text: str) -> Dict[str, str]:
        contact = {'name': 'Unknown', 'email': '', 'phone': ''}
        
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact['email'] = email_match.group()
        
        phone_match = re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        lines = text.split('\n')
        for line in lines[:3]:
            if line.strip() and len(line.strip()) < 50 and '@' not in line:
                contact['name'] = line.strip()
                break
        
        return contact
    
    def _extract_experience(self, text: str) -> int:
        patterns = [r'(\d+)\+?\s*years?\s*(?:of\s*)?experience', r'experience:\s*(\d+)\+?\s*years?']
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return int(match.group(1))
        return 0
    
    def _detect_bias(self, text: str) -> Dict[str, Any]:
        bias_indicators = ['he', 'she', 'him', 'his', 'her', 'years old', 'married', 'single', 'children']
        text_lower = text.lower()
        detected = [item for item in bias_indicators if item in text_lower]
        
        risk_score = min(len(detected) * 10, 100)
        risk_level = 'low' if risk_score <= 33 else 'medium' if risk_score <= 66 else 'high'
        
        return {
            'overall_risk_score': risk_score,
            'overall_risk_level': risk_level,
            'detected_items': detected,
            'recommendations': ['Remove personal information'] if detected else []
        }
    
    def _calculate_skill_match(self, job_skills: Dict[str, List[str]], resume_skills: Dict[str, List[str]]) -> Dict[str, Any]:
        total_job_skills = sum(len(skills) for skills in job_skills.values())
        if total_job_skills == 0:
            return {'overall_coverage': 0, 'matched_skills': [], 'missing_skills': []}
        
        matched = []
        for category, skills in job_skills.items():
            resume_cat_skills = resume_skills.get(category, [])
            matched.extend(set(skills) & set(resume_cat_skills))
        
        overall_coverage = len(set(matched)) / total_job_skills * 100
        
        return {
            'overall_coverage': overall_coverage,
            'matched_skills': list(set(matched)),
            'missing_skills': []
        }
    
    def _keyword_similarity(self, text1: str, text2: str) -> float:
        words1 = set(re.findall(r'\b\w+\b', text1.lower()))
        words2 = set(re.findall(r'\b\w+\b', text2.lower()))
        
        if not words2:
            return 50.0
        
        overlap = len(words1 & words2)
        return (overlap / len(words2)) * 100


class StandaloneVisualizer:
    """Fallback visualizer."""
    
    def create_ranking_bar_chart(self, candidates: List[Dict[str, Any]]):
        import plotly.express as px
        import plotly.graph_objects as go
        
        names = [c.get('name', f"Candidate {i+1}") for i, c in enumerate(candidates)]
        scores = [c.get('final_score', 0) for c in candidates]
        
        colors = ['#2ca02c' if s >= 70 else '#ff7f0e' if s >= 50 else '#d62728' for s in scores]
        
        fig = go.Figure(data=[
            go.Bar(y=names, x=scores, orientation='h', marker_color=colors,
                   text=[f"{s:.1f}" for s in scores], textposition='auto')
        ])
        
        fig.update_layout(xaxis_title="Final Score", 
                         yaxis_title="Candidates", height=max(400, len(candidates) * 40))
        fig.update_yaxes(autorange="reversed")
        
        return fig
    
    def create_ranking_chart(self, candidates: List[Dict[str, Any]]):
        """Create ranking chart (alias for create_ranking_bar_chart)."""
        return self.create_ranking_bar_chart(candidates)
    
    def create_experience_distribution_chart(self, data: List[Dict[str, Any]]):
        """Create experience distribution chart (used for bias distribution)."""
        import plotly.express as px
        import plotly.graph_objects as go
        
        # Extract years_of_experience values
        values = [item.get('years_of_experience', 0) for item in data]
        
        fig = go.Figure(data=[
            go.Histogram(
                x=values,
                nbinsx=10,
                marker_color='#1f77b4',
                opacity=0.75,
                hovertemplate='Experience: %{x} years<br>Count: %{y}<extra></extra>'
            )
        ])
        
        fig.update_layout(
            title="Experience Distribution",
            xaxis_title="Years of Experience",
            yaxis_title="Number of Candidates",
            height=400,
            showlegend=False
        )
        
        return fig
    
    def create_candidate_comparison_chart(self, candidate1: Dict[str, Any], candidate2: Dict[str, Any]):
        """Create comparison chart for two candidates."""
        import plotly.graph_objects as go
        
        # Metrics to compare
        metrics = ['Final Score', 'Semantic Similarity', 'Skill Coverage', 'Experience']
        
        # Extract values for candidate 1
        values1 = [
            candidate1.get('final_score', 0),
            candidate1.get('semantic_similarity', 0),
            candidate1.get('skill_coverage', {}).get('overall_coverage', 0),
            candidate1.get('years_of_experience', 0)
        ]
        
        # Extract values for candidate 2
        values2 = [
            candidate2.get('final_score', 0),
            candidate2.get('semantic_similarity', 0),
            candidate2.get('skill_coverage', {}).get('overall_coverage', 0),
            candidate2.get('years_of_experience', 0)
        ]
        
        # Create figure
        fig = go.Figure()
        
        # Add candidate 1
        fig.add_trace(go.Bar(
            name=candidate1.get('name', 'Candidate 1'),
            x=metrics,
            y=values1,
            marker_color='#1f77b4'
        ))
        
        # Add candidate 2
        fig.add_trace(go.Bar(
            name=candidate2.get('name', 'Candidate 2'),
            x=metrics,
            y=values2,
            marker_color='#ff7f0e'
        ))
        
        # Update layout
        fig.update_layout(
            title="Candidate Comparison",
            xaxis_title="Metrics",
            yaxis_title="Score",
            height=400,
            barmode='group',
            showlegend=True
        )
        
        return fig


# Configure Streamlit page
st.set_page_config(
    page_title="HireSense AI - Intelligent Resume Screening and Ranking System",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for better styling with dark mode support
st.markdown("""
<style>
    /* Hide sidebar by default and only show when expanded */
    [data-testid="stSidebar"] {
        display: none !important;
    }
    
    [data-testid="stSidebar"][aria-expanded="true"] {
        display: block !important;
    }
    
    .main-header {
        font-size: 2.2rem;
        color: #ffffff;
        text-align: center;
        padding: 1.5rem 1rem;
        background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%);
        border-radius: 15px;
        box-shadow: 0 8px 32px rgba(220, 53, 69, 0.3);
        text-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
        margin: 1rem 0 2rem 0;
    }
    .metric-card {
        background-color: #f8f9fa;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
        margin: 0.5rem 0;
    }
    .candidate-card {
        background-color: var(--background-color, white);
        color: var(--text-color, #333333);
        padding: 1.5rem;
        border-radius: 0.5rem;
        border: 1px solid var(--border-color, #dee2e6);
        margin: 1rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .score-high { color: #2ca02c; font-weight: bold; }
    .score-medium { color: #ff7f0e; font-weight: bold; }
    .score-low { color: #d62728; font-weight: bold; }
    
    /* Dark mode support */
    @media (prefers-color-scheme: dark) {
        .main-header {
            color: #ffffff;
            background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%);
        }
        .metric-card {
            background-color: #1e1e1e;
            border-left-color: #4fc3f7;
        }
        .candidate-card {
            background-color: #1f2937;
            color: #ffffff;
            border-color: #4a5568;
        }
    }
    
    /* Streamlit dark mode support */
    [data-testid="stAppViewContainer"] {
        background-color: var(--background-color);
        color: var(--text-color);
    }
    
    /* Landing page styles */
    .landing-container {
        display: flex;
        flex-direction: column;
        justify-content: flex-start;
        align-items: center;
        min-height: 40vh;
        text-align: center;
        padding: 1rem;
        padding-top: 6rem;
    }
    
    .landing-title {
        font-size: 3rem;
        font-weight: 700;
        color: #ffffff;
        margin-bottom: 0.5rem;
        text-align: center;
        animation: fadeInUp 1s ease-out;
        background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        text-shadow: 0 2px 8px rgba(220, 53, 69, 0.3);
    }
    
    .landing-subtitle {
        font-size: 1.2rem;
        color: #b0b0b0;
        margin-bottom: 0.3rem;
        text-align: center;
        animation: fadeInUp 1s ease-out 0.2s both;
    }
    
    .landing-enter-btn {
        background: linear-gradient(135deg, #1f77b4, #17a2b8);
        color: white;
        border: none;
        padding: 0.8rem 2.5rem;
        font-size: 1.2rem;
        font-weight: 600;
        border-radius: 50px;
        cursor: pointer;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(31, 119, 180, 0.3);
        animation: fadeInUp 1s ease-out 0.4s both;
    }
    
    .landing-enter-btn:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(31, 119, 180, 0.4);
        background: linear-gradient(135deg, #17a2b8, #1f77b4);
    }
    
    @keyframes fadeInUp {
        from {
            opacity: 0;
            transform: translateY(30px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    /* Dark mode for landing page */
    @media (prefers-color-scheme: dark) {
        .landing-title {
            color: #4fc3f7;
        }
        .landing-subtitle {
            color: #b0b0b0;
        }
    }
    </style>
""", unsafe_allow_html=True)


class ResumeRankingApp:
    """Main application class for the resume ranking system."""
    
    def __init__(self):
        """Initialize the application."""
        # Always use fallback mode to avoid torch issues
        self.scoring_engine = StandaloneRanker()
        self.visualizer = StandaloneVisualizer()
        self.text_extractor = None  # Not used in fallback mode
        
        # Initialize session state
        if 'job_analysis' not in st.session_state:
            st.session_state.job_analysis = None
        if 'candidates' not in st.session_state:
            st.session_state.candidates = []
        if 'current_page' not in st.session_state:
            st.session_state.current_page = "Home"
        if 'show_landing' not in st.session_state:
            st.session_state.show_landing = True
    
    def run(self):
        """Run the main application."""
        
        # Show landing page if first visit
        if st.session_state.show_landing:
            self.render_landing_page()
            return
        
        # Show main app
        self.render_sidebar()
        self.render_header()
        
        # Page routing
        if st.session_state.current_page == "Home":
            self.render_home()
        elif st.session_state.current_page == "Upload Resumes":
            self.render_upload_page()
        elif st.session_state.current_page == "Ranking Dashboard":
            self.render_ranking_dashboard()
        elif st.session_state.current_page == "Candidate Insights":
            self.render_candidate_insights()
        elif st.session_state.current_page == "Bias Analysis":
            self.render_bias_analysis()
        elif st.session_state.current_page == "Resume Comparison":
            self.render_resume_comparison()
    
    def render_sidebar(self):
        """Render the sidebar with minimal spacing."""
        with st.sidebar:
            st.title("⚙️ Settings")
            st.markdown("<div style='border-bottom: 1px solid #f0f0f0; margin-bottom: 10px;'></div>", unsafe_allow_html=True)
            
            st.subheader("📉 System Sensitivity")
            
            # Configuration
            with st.expander("🔧 Configuration", expanded=False):
                # Scoring Weights
                st.write("**📊 Scoring Weights**")
                skill_weight = st.slider("Skill Match", 0, 100, 60, 5)
                exp_weight = st.slider("Experience", 0, 100, 25, 5)
                semantic_weight = st.slider("Semantic Match", 0, 100, 15, 5)
                
                total_weight = skill_weight + exp_weight + semantic_weight
                if total_weight != 100:
                    st.warning(f"⚠️ Total: {total_weight}% - Auto-normalizing")
                    skill_weight = int(skill_weight * 100 / total_weight)
                    exp_weight = int(exp_weight * 100 / total_weight)
                    semantic_weight = 100 - skill_weight - exp_weight
                
                st.session_state.scoring_weights = {
                    'skill_match': skill_weight,
                    'experience': exp_weight,
                    'semantic_similarity': semantic_weight
                }
                
                # Bias Detection
                st.markdown("<div style='background-color: #1a1a1a; padding: 5px; border-radius: 5px; margin-bottom: 10px;'>", unsafe_allow_html=True)
                st.write("**⚖️ Bias Detection**")
                enable_bias_detection = st.checkbox("Enable Detection", value=True)
                bias_sensitivity = st.selectbox("Sensitivity", ["Low", "Medium", "High"], index=1)
                
                st.session_state.bias_settings = {
                    'enabled': enable_bias_detection,
                    'sensitivity': bias_sensitivity
                }
                st.markdown("</div>", unsafe_allow_html=True)
                
                # Display
                st.markdown("<div style='background-color: #1a1a1a; padding: 5px; border-radius: 5px; margin-bottom: 10px;'>", unsafe_allow_html=True)
                st.write("**🎨 Display**")
                show_explanations = st.checkbox("Detailed Explanations", value=True)
                show_charts = st.checkbox("Interactive Charts", value=True)
                
                st.session_state.display_settings = {
                    'show_explanations': show_explanations,
                    'show_charts': show_charts
                }
                st.markdown("</div>", unsafe_allow_html=True)
            
            st.markdown("---")
            
            # System Status
            st.subheader("📊 System Status")
            
            if st.session_state.job_analysis:
                st.success("✅ Job Description Analyzed")
            else:
                st.warning("⚠️ No job description")
            
            if st.session_state.candidates:
                st.success(f"✅ {len(st.session_state.candidates)} Candidate's Resumes Analyzed")
                avg_score = sum(c.get('final_score', 0) for c in st.session_state.candidates) / len(st.session_state.candidates)
                st.write(f"📈 Avg Score: {avg_score:.1f}/100")
                
                top_candidate = max(st.session_state.candidates, key=lambda x: x.get('final_score', 0))
                st.write(f"🏆 Top Candidate: {top_candidate.get('name', 'Unknown')}")
            else:
                st.info("ℹ️ No resumes uploaded")
            
            st.markdown("---")
            
            # Quick Actions
            st.subheader("⚡ Quick Actions")
            
            if st.button("🔄 Refresh", use_container_width=True):
                # Clear all session state for complete refresh
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                # Reinitialize essential session state variables
                st.session_state.job_analysis = None
                st.session_state.candidates = []
                st.session_state.current_page = "Home"
                st.session_state.show_landing = False
                st.success("✅ System refreshed!")
                st.rerun()
            
            st.markdown("---")
            
            # Powered by section
            st.markdown(
                """
                <div style='text-align: center; color: #666; font-size: 0.8rem; padding: 0.5rem;'>
                    <strong>🚀 Powered by Streamlit</strong>
                </div>
                """, 
                unsafe_allow_html=True
            )
    
    
    
    def render_header(self):
        """Render the main header."""
        st.markdown('''
        <div class="main-header">
            <h1 style="margin: 0; font-size: 2.2rem;">HireSense AI 🧠</h1>
            <h2 style="margin: -0.3rem 0 0 0; font-size: 1.2rem; font-weight: 400;">Intelligent Resume Screening & Ranking System</h2>
        </div>
        ''', unsafe_allow_html=True)
    
    def render_landing_page(self):
        """Render the landing page with centered content."""
        # Hide sidebar and main content for landing page
        st.markdown("""
        <style>
            [data-testid="stSidebar"] {
                display: none;
            }
            .stMain {
                padding-top: 0;
            }
        </style>
        """, unsafe_allow_html=True)
        
        # Create centered landing page
        st.markdown("""
        <div class="landing-container">
            <h1 class="landing-title">HireSense AI 🧠</h1>
            <p class="landing-subtitle">An Intelligent Resume Screening and Ranking System</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Center the enter button with reduced width
        col1, col2, col3 = st.columns([3, 2, 3])
        with col2:
            if st.button("Enter➡️", key="landing_enter_btn", use_container_width=True):
                st.session_state.show_landing = False
                st.rerun()
    
    def render_home(self):
        """Render main dashboard with all features."""
        st.markdown("---")
        
        # Job Description Upload - Always visible
        st.subheader("📋 Job Description Analysis")
        job_text = st.text_area(
            "Paste the job description here:",
            height=150,
            placeholder="Enter the complete job description including requirements, responsibilities, and qualifications..."
        )
        
        if st.button("🔍 Analyze Job Description", type="primary"):
            if job_text.strip():
                with st.spinner("Analyzing job description..."):
                    try:
                        # Use fallback method
                        job_analysis = self.scoring_engine.analyze_job(job_text)
                        
                        st.session_state.job_analysis = job_analysis
                        st.success("✅ Job description analyzed successfully!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error analyzing job description: {e}")
            else:
                st.error("Please enter a job description")
        
        # Display job analysis results if available
        if st.session_state.job_analysis:
            job_analysis = st.session_state.job_analysis
            
            # Show enhanced metrics display
            st.markdown("""
            <div style="background: linear-gradient(135deg, #dc3545 0%, #c62828 100%); 
                        padding: 0.5rem; border-radius: 8px; margin: 0.3rem 0; 
                        box-shadow: 0 1px 4px rgba(220,53,69,0.3);">
                <div style="display: flex; justify-content: space-around; align-items: center;">
                    <div style="text-align: center; color: white;">
                        <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                            📊
                        </div>
                        <div style="font-size: 1.1rem; font-weight: bold;">
                            {}
                        </div>
                        <div style="font-size: 0.7rem; opacity: 0.9;">
                            Skill Categories
                        </div>
                    </div>
                    <div style="text-align: center; color: white;">
                        <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                            ⚖️
                        </div>
                        <div style="font-size: 1.1rem; font-weight: bold; text-transform: capitalize;">
                            {}
                        </div>
                        <div style="font-size: 0.7rem; opacity: 0.9;">
                            Bias Risk Level
                        </div>
                    </div>
                    <div style="text-align: center; color: white;">
                        <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                            💼
                        </div>
                        <div style="font-size: 1.1rem; font-weight: bold;">
                            {}+ years
                        </div>
                        <div style="font-size: 0.7rem; opacity: 0.9;">
                            Experience Required
                        </div>
                    </div>
                </div>
            </div>
            """.format(
                len(job_analysis['skills']),
                job_analysis['bias_analysis'].get('overall_risk_level' if 'overall_risk_level' in job_analysis['bias_analysis'] else 'risk_level', 'N/A'),
                job_analysis['experience_requirements'].get('minimum_years', 0)
            ), unsafe_allow_html=True)
            
            # Show extracted skills
            if job_analysis['skills']:
                st.subheader("🎯 Required Skills by Category")
                
                # Create tabular format
                skills_data = []
                for category, skill_list in job_analysis['skills'].items():
                    if skill_list:
                        skills_data.append({
                            'Skill Category': category,
                            'Required Skills': ', '.join(skill_list)
                        })
                
                if skills_data:
                    skills_df = pd.DataFrame(skills_data)
                    st.dataframe(skills_df, use_container_width=True, hide_index=True)
        
        # Resume Upload - Visible after job analysis
        if st.session_state.job_analysis:
            st.markdown("---")
            st.subheader("📄 Resume Upload")
            
            # Upload method selection
            upload_method = st.radio(
                "Choose upload method:",
                ["📝 Paste Resume Texts", "📁 Upload Resume Files", "📊 Upload Dataset"],
                horizontal=True
            )
            
            if upload_method == "📁 Upload Resume Files":
                # File upload mode
                st.info("📂 Upload PDF/DOCX files (multiple files supported)")
                
                # Check if required libraries are available
                if not PDF_AVAILABLE or not DOCX_AVAILABLE:
                    st.error("❌ File upload requires PyPDF2 and python-docx libraries. Please install them:")
                    st.write("- `pip install PyPDF2`")
                    st.write("- `pip install python-docx`")
                    st.write("Or use the '📝 Paste Resume Texts' option instead.")
                    return
                
                uploaded_files = st.file_uploader(
                    "Choose resume files:",
                    type=['pdf', 'docx'],
                    accept_multiple_files=True,
                    help="Upload multiple resume files in PDF or DOCX format"
                )
                
                if uploaded_files:
                    if st.button("🚀 Analyze Uploaded Files", type="primary"):
                        with st.spinner("Processing uploaded files..."):
                            try:
                                candidates = []
                                for i, file in enumerate(uploaded_files):
                                    # Extract text from file
                                    if file.type == "application/pdf":
                                        import PyPDF2
                                        reader = PyPDF2.PdfReader(file)
                                        text = ""
                                        for page in reader.pages:
                                            text += page.extract_text()
                                    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                        import docx
                                        doc = docx.Document(file)
                                        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                                    else:
                                        text = str(file.read())
                                    
                                    # Analyze resume
                                    candidate = self.scoring_engine.analyze_resume(
                                        text,
                                        file.name,
                                        st.session_state.job_analysis,
                                        st.session_state.get('scoring_weights', {
                                            'skill_match': 60,
                                            'experience': 25,
                                            'semantic_similarity': 15
                                        })
                                    )
                                    candidates.append(candidate)
                                
                                # Rank candidates before storing
                                ranked_candidates = self.scoring_engine.rank_candidates(
                                    candidates, st.session_state.job_analysis
                                )
                                st.session_state.candidates = ranked_candidates
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error processing files: {e}")
                                st.markdown("""
                                <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                                            padding: 1.5rem; border-radius: 15px; margin: 1rem 0; 
                                            box-shadow: 0 4px 15px rgba(0,0,0,0.1);">
                                    <div style="display: flex; justify-content: space-around; align-items: center;">
                                        <div style="text-align: center; color: white;">
                                            <div style="font-size: 2rem; font-weight: bold; margin-bottom: 0.5rem;">
                                                📊
                                            </div>
                                            <div style="font-size: 1.8rem; font-weight: bold;">
                                                {}
                                            </div>
                                            <div style="font-size: 0.9rem; opacity: 0.9;">
                                                Skill Categories
                                            </div>
                                        </div>
                                        <div style="text-align: center; color: white;">
                                            <div style="font-size: 2rem; font-weight: bold; margin-bottom: 0.5rem;">
                                                ⚖️
                                            </div>
                                            <div style="font-size: 1.8rem; font-weight: bold; text-transform: capitalize;">
                                                {}
                                            </div>
                                            <div style="font-size: 0.9rem; opacity: 0.9;">
                                                Bias Risk Level
                                            </div>
                                        </div>
                                        <div style="text-align: center; color: white;">
                                            <div style="font-size: 2rem; font-weight: bold; margin-bottom: 0.5rem;">
                                                💼
                                            </div>
                                            <div style="font-size: 1.8rem; font-weight: bold;">
                                                {}+ years
                                            </div>
                                            <div style="font-size: 0.9rem; opacity: 0.9;">
                                                Experience Required
                                            </div>
                                        </div>
                                    </div>
                                </div>
                                """.format(
                                    len(st.session_state.job_analysis['skills']),
                                    st.session_state.job_analysis['bias_analysis'].get('overall_risk_level' if 'overall_risk_level' in st.session_state.job_analysis['bias_analysis'] else 'risk_level', 'N/A'),
                                    st.session_state.job_analysis['experience_requirements'].get('minimum_years', 0)
                                ), unsafe_allow_html=True)
                                
                            except Exception as e:
                                st.error(f"Error processing files: {e}")
            
            elif upload_method == "📊 Upload Dataset":
                # Dataset upload mode
                st.info("📂 Upload a dataset file containing multiple resumes (CSV, JSON, or ZIP with PDF/DOCX files)")
                
                dataset_file = st.file_uploader(
                    "Upload dataset file:",
                    type=['csv', 'json', 'zip'],
                    accept_multiple_files=False,
                    help="Upload a CSV/JSON file with resume data or a ZIP file containing multiple resume files"
                )
                
                if dataset_file and st.button("🚀 Analyze Dataset", type="primary"):
                    with st.spinner("Processing dataset..."):
                        try:
                            import zipfile
                            import tempfile
                            import os
                            
                            candidates = []
                            
                            # Debug: Show file information
                            st.write(f"📁 File name: {dataset_file.name}")
                            st.write(f"📄 File type: {dataset_file.type}")
                            st.write(f"📋 File size: {dataset_file.size} bytes")
                            
                            # Handle ZIP file with multiple resume files
                            if dataset_file.type == "application/zip" or dataset_file.name.lower().endswith('.zip'):
                                st.write("🔍 Processing as ZIP file...")
                                
                                with tempfile.TemporaryDirectory() as temp_dir:
                                    with zipfile.ZipFile(dataset_file) as zip_file:
                                        zip_file.extractall(temp_dir)
                                    
                                    # Debug: Show extracted structure
                                    st.write(f"📁 Extracted to: {temp_dir}")
                                    
                                    # Recursively find all resume files in subdirectories
                                    resume_files = []
                                    all_files = []
                                    for root, dirs, files in os.walk(temp_dir):
                                        for filename in files:
                                            file_path = os.path.join(root, filename)
                                            all_files.append(file_path)
                                            if filename.lower().endswith(('.pdf', '.docx')):
                                                resume_files.append(file_path)
                                    
                                    # Debug: Show all files found
                                    st.write(f"📄 Total files found: {len(all_files)}")
                                    if all_files:
                                        st.write("📋 Files found:")
                                        for f in all_files:
                                            rel_path = os.path.relpath(f, temp_dir)
                                            st.write(f"  • {rel_path}")
                                    else:
                                        st.write("❌ No files found in ZIP")
                                    
                                    # Debug: Show resume files found
                                    st.write(f"📄 Resume files found: {len(resume_files)}")
                                    if resume_files:
                                        st.write("📋 Resume files to process:")
                                        for f in resume_files:
                                            rel_path = os.path.relpath(f, temp_dir)
                                            st.write(f"  • {rel_path}")
                                    
                                    # Process all found resume files
                                    processed_count = 0
                                    for file_path in resume_files:
                                        filename = os.path.basename(file_path)
                                        st.write(f"🔍 Processing: {filename}")
                                        
                                        try:
                                            if filename.lower().endswith('.pdf'):
                                                if PDF_AVAILABLE:
                                                    import PyPDF2
                                                    reader = PyPDF2.PdfReader(file_path)
                                                    text = ""
                                                    for page in reader.pages:
                                                        text += page.extract_text()
                                                    
                                                    if text.strip():
                                                        candidate = self.scoring_engine.analyze_resume(
                                                            text, filename, st.session_state.job_analysis
                                                        )
                                                        candidates.append(candidate)
                                                        processed_count += 1
                                                        st.write(f"✅ Successfully processed: {filename}")
                                                    else:
                                                        st.write(f"⚠️ Empty PDF: {filename}")
                                            
                                            elif filename.lower().endswith('.docx'):
                                                if DOCX_AVAILABLE:
                                                    import docx
                                                    doc = docx.Document(file_path)
                                                    text = "\n".join([para.text for para in doc.paragraphs])
                                                    
                                                    if text.strip():
                                                        candidate = self.scoring_engine.analyze_resume(
                                                            text, filename, st.session_state.job_analysis
                                                        )
                                                        candidates.append(candidate)
                                                        processed_count += 1
                                                        st.write(f"✅ Successfully processed: {filename}")
                                                    else:
                                                        st.write(f"⚠️ Empty DOCX: {filename}")
                                        
                                        except Exception as e:
                                            st.write(f"❌ Error processing {filename}: {e}")
                                    
                                    st.write(f"📊 Summary: {processed_count} files processed successfully")
                            
                            elif dataset_file.type == "text/csv":
                                st.write("🔍 Processing as CSV file...")
                                # Handle CSV file with resume data
                                df = pd.read_csv(dataset_file)
                                
                                # Look for common column names for resume text
                                text_column = None
                                for col in ['resume_text', 'text', 'content', 'resume', 'description']:
                                    if col in df.columns:
                                        text_column = col
                                        break
                                
                                if text_column:
                                    for idx, row in df.iterrows():
                                        resume_text = str(row[text_column])
                                        if resume_text.strip():
                                            candidate_name = str(row.get('name', f'Candidate_{idx+1}'))
                                            candidate = self.scoring_engine.analyze_resume(
                                                resume_text, candidate_name, st.session_state.job_analysis
                                            )
                                            candidates.append(candidate)
                                else:
                                    st.error("❌ CSV file must contain a 'resume_text', 'text', 'content', 'resume', or 'description' column")
                                    return
                            
                            elif dataset_file.type == "application/json":
                                st.write("🔍 Processing as JSON file...")
                                # Handle JSON file with resume data
                                json_data = json.loads(dataset_file.read().decode('utf-8'))
                                
                                if isinstance(json_data, list):
                                    for idx, item in enumerate(json_data):
                                        if isinstance(item, dict):
                                            # Look for resume text in various fields
                                            resume_text = ""
                                            for field in ['resume_text', 'text', 'content', 'resume', 'description']:
                                                if field in item:
                                                    resume_text = str(item[field])
                                                    break
                                            
                                            if resume_text.strip():
                                                candidate_name = str(item.get('name', f'Candidate_{idx+1}'))
                                                candidate = self.scoring_engine.analyze_resume(
                                                    resume_text, candidate_name, st.session_state.job_analysis
                                                )
                                                candidates.append(candidate)
                                else:
                                    st.error("❌ JSON file must contain an array of resume objects")
                                    return
                            
                            else:
                                st.error(f"❌ Unsupported file type: {dataset_file.type}")
                                st.write("Supported formats: ZIP, CSV, JSON")
                                return
                            
                            # Rank candidates
                            if candidates:
                                ranked_candidates = self.scoring_engine.rank_candidates(
                                    candidates, st.session_state.job_analysis
                                )
                                st.session_state.candidates = ranked_candidates
                                
                                st.success(f"✅ Successfully processed {len(candidates)} resumes from dataset!")
                                st.rerun()
                            else:
                                st.warning("⚠️ No valid resumes found in dataset")
                                
                        except Exception as e:
                            st.error(f"Error processing dataset: {e}")
            
            else:
                # Fallback text input mode
                st.info("📝 Paste resume texts below (separate with '---')")
                resume_texts = st.text_area(
                    "Paste resume texts (one per resume, separated by '---'):",
                    height=200,
                    placeholder="Paste resume content here. Separate multiple resumes with '---' on a new line."
                )
                
                if st.button("🚀 Analyze Resumes", type="primary"):
                    if resume_texts.strip():
                        with st.spinner("Processing resumes..."):
                            try:
                                # Split resumes by separator
                                resume_list = [text.strip() for text in resume_texts.split('---') if text.strip()]
                                
                                candidates = []
                                for i, resume_text in enumerate(resume_list):
                                    # Analyze resume
                                    candidate = self.scoring_engine.analyze_resume(
                                        resume_text,
                                        f"Resume_{i+1}",
                                        st.session_state.job_analysis
                                    )
                                    candidates.append(candidate)
                                
                                # Rank candidates
                                if candidates:
                                    ranked_candidates = self.scoring_engine.rank_candidates(
                                        candidates, st.session_state.job_analysis
                                    )
                                    st.session_state.candidates.extend(ranked_candidates)
                            except Exception as e:
                                st.error(f"Error processing resumes: {e}")
                    else:
                        st.error("Please enter resume texts")
        
        # Analysis Selection - Visible after resumes are uploaded
        if st.session_state.candidates:
            st.success(f"✅ Successfully analyzed {len(st.session_state.candidates)} resumes! Refer below for Detailed Rankings ")
            st.divider()
            st.subheader("🕵️ Detailed Analysis")
            
            # Initialize analysis selection state
            if 'selected_analysis' not in st.session_state:
                st.session_state.selected_analysis = None
            
            # Horizontal selection buttons
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("📊 Ranking Dashboard", type="primary" if st.session_state.selected_analysis == "ranking" else "secondary", use_container_width=True):
                    st.session_state.selected_analysis = "ranking"
                    st.rerun()
            
            with col2:
                if st.button("📈 Candidate Insights", type="primary" if st.session_state.selected_analysis == "insights" else "secondary", use_container_width=True):
                    st.session_state.selected_analysis = "insights"
                    st.rerun()
            
            with col3:
                if st.button("⚖️ Bias Analysis", type="primary" if st.session_state.selected_analysis == "bias" else "secondary", use_container_width=True):
                    st.session_state.selected_analysis = "bias"
                    st.rerun()
            
            with col4:
                if st.button("🔄 Resume Comparison", type="primary" if st.session_state.selected_analysis == "comparison" else "secondary", use_container_width=True):
                    st.session_state.selected_analysis = "comparison"
                    st.rerun()
            
            st.divider()
            
            # Display selected analysis
            if st.session_state.selected_analysis == "ranking":
                self.render_ranking_dashboard()
            elif st.session_state.selected_analysis == "insights":
                self.render_candidate_insights()
            elif st.session_state.selected_analysis == "bias":
                self.render_bias_analysis()
            elif st.session_state.selected_analysis == "comparison":
                self.render_resume_comparison()
        
                
                
    
    def render_upload_page(self):
        """Render the upload page."""
        st.header("📄 Upload Job Description & Resumes")
        
        # Job description upload
        st.subheader("📋 Job Description")
        job_text = st.text_area(
            "Paste the job description here:",
            height=200,
            placeholder="Enter the complete job description including requirements, responsibilities, and qualifications..."
        )
        
        if st.button("🔍 Analyze Job Description", type="primary"):
            if job_text.strip():
                with st.spinner("Analyzing job description..."):
                    try:
                        # Use fallback method
                        job_analysis = self.scoring_engine.analyze_job(job_text)
                        
                        st.session_state.job_analysis = job_analysis
                        
                        # Show job analysis summary
                        st.success("✅ Job description analyzed successfully!")
                        
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.metric("Skill Categories", len(job_analysis['skills']))
                        with col2:
                            bias_key = 'overall_risk_level' if 'overall_risk_level' in job_analysis['bias_analysis'] else 'risk_level'
                            st.metric("Bias Risk Level", job_analysis['bias_analysis'][bias_key])
                        with col3:
                            exp_key = 'minimum_years' if 'minimum_years' in job_analysis['experience_requirements'] else 'minimum_years'
                            exp_years = job_analysis['experience_requirements'].get(exp_key, 0)
                            st.metric("Experience Required", f"{exp_years}+ years")
                        
                        # Show extracted skills
                        if job_analysis['skills']:
                            st.subheader("🎯 Required Skills by Category")
                            for category, skills in job_analysis['skills'].items():
                                if skills:
                                    st.write(f"**{category}**: {', '.join(skills)}")
                    
                    except Exception as e:
                        st.error(f"Error analyzing job description: {e}")
            else:
                st.error("Please enter a job description")
        
        st.divider()
        
        # Resume upload
        st.subheader("📁 Upload Resumes")
        
        if not st.session_state.job_analysis:
            st.warning("⚠️ Please analyze a job description first")
            return
        
        # Upload method selection
        upload_method = st.radio(
            "Choose upload method:",
            ["📝 Paste Resume Texts", "📁 Upload Resume Files", "📊 Upload Dataset"],
            horizontal=True
        )
        
        if upload_method == "📁 Upload Resume Files":
            # File upload mode
            st.info("📂 Upload PDF/DOCX files (multiple files supported)")
            
            # Check if required libraries are available
            if not PDF_AVAILABLE or not DOCX_AVAILABLE:
                st.error("❌ File upload requires PyPDF2 and python-docx libraries. Please install them:")
                st.write("- `pip install PyPDF2`")
                st.write("- `pip install python-docx`")
                st.write("Or use the '📝 Paste Resume Texts' option instead.")
                return
            
            uploaded_files = st.file_uploader(
                "Choose resume files:",
                type=['pdf', 'docx'],
                accept_multiple_files=True,
                help="Upload multiple resume files in PDF or DOCX format"
            )
            
            if uploaded_files:
                if st.button("🚀 Analyze Uploaded Files", type="primary"):
                    with st.spinner("Processing uploaded files..."):
                        try:
                            candidates = []
                            for i, file in enumerate(uploaded_files):
                                # Extract text from file
                                if file.type == "application/pdf":
                                    import PyPDF2
                                    reader = PyPDF2.PdfReader(file)
                                    text = ""
                                    for page in reader.pages:
                                        text += page.extract_text()
                                elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                    import docx
                                    doc = docx.Document(file)
                                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                                else:
                                    text = str(file.read())
                                
                                # Analyze resume
                                candidate = self.scoring_engine.analyze_resume(
                                    text,
                                    file.name,
                                    st.session_state.job_analysis
                                )
                                candidates.append(candidate)
                            
                            # Rank candidates
                            if candidates:
                                ranked_candidates = self.scoring_engine.rank_candidates(
                                    candidates, st.session_state.job_analysis
                                )
                                st.session_state.candidates.extend(ranked_candidates)
                                st.success(f"✅ Successfully analyzed {len(candidates)} resume files!")
                                st.rerun()
                            else:
                                st.error("No resumes could be processed successfully")
                        
                        except Exception as e:
                            st.error(f"Error processing files: {e}")
        
        elif upload_method == "📊 Upload Dataset":
            # Dataset upload mode
            st.info("📂 Upload a dataset file containing multiple resumes (CSV, JSON, or ZIP with PDF/DOCX files)")
            
            dataset_file = st.file_uploader(
                "Upload dataset file:",
                type=['csv', 'json', 'zip'],
                accept_multiple_files=False,
                help="Upload a CSV/JSON file with resume data or a ZIP file containing multiple resume files"
            )
            
            if dataset_file and st.button("🚀 Analyze Dataset", type="primary"):
                with st.spinner("Processing dataset..."):
                    try:
                        import zipfile
                        import tempfile
                        import os
                        
                        candidates = []
                        
                        if dataset_file.type == "application/zip":
                            # Handle ZIP file with multiple resume files
                            with tempfile.TemporaryDirectory() as temp_dir:
                                with zipfile.ZipFile(dataset_file) as zip_file:
                                    zip_file.extractall(temp_dir)
                                
                                # Debug: Show extracted structure
                                st.write(f"📁 Extracted to: {temp_dir}")
                                
                                # Recursively find all resume files in subdirectories
                                resume_files = []
                                all_files = []
                                for root, dirs, files in os.walk(temp_dir):
                                    for filename in files:
                                        file_path = os.path.join(root, filename)
                                        all_files.append(file_path)
                                        if filename.lower().endswith(('.pdf', '.docx')):
                                            resume_files.append(file_path)
                                
                                # Debug: Show all files found
                                st.write(f"📄 Total files found: {len(all_files)}")
                                if all_files:
                                    st.write("📋 Files found:")
                                    for f in all_files:
                                        rel_path = os.path.relpath(f, temp_dir)
                                        st.write(f"  • {rel_path}")
                                else:
                                    st.write("❌ No files found in ZIP")
                                
                                # Debug: Show resume files found
                                st.write(f"📄 Resume files found: {len(resume_files)}")
                                if resume_files:
                                    st.write("📋 Resume files to process:")
                                    for f in resume_files:
                                        rel_path = os.path.relpath(f, temp_dir)
                                        st.write(f"  • {rel_path}")
                                
                                # Process all found resume files
                                for file_path in resume_files:
                                    filename = os.path.basename(file_path)
                                    
                                    if filename.lower().endswith('.pdf'):
                                        if PDF_AVAILABLE:
                                            import PyPDF2
                                            reader = PyPDF2.PdfReader(file_path)
                                            text = ""
                                            for page in reader.pages:
                                                text += page.extract_text()
                                            
                                            if text.strip():
                                                candidate = self.scoring_engine.analyze_resume(
                                                    text, filename, st.session_state.job_analysis
                                                )
                                                candidates.append(candidate)
                                    
                                    elif filename.lower().endswith('.docx'):
                                        if DOCX_AVAILABLE:
                                            import docx
                                            doc = docx.Document(file_path)
                                            text = "\n".join([para.text for para in doc.paragraphs])
                                            
                                            if text.strip():
                                                candidate = self.scoring_engine.analyze_resume(
                                                    text, filename, st.session_state.job_analysis
                                                )
                                                candidates.append(candidate)
                        
                        elif dataset_file.type == "text/csv":
                            # Handle CSV file with resume data
                            df = pd.read_csv(dataset_file)
                            
                            # Look for common column names for resume text
                            text_column = None
                            for col in ['resume_text', 'text', 'content', 'resume', 'description']:
                                if col in df.columns:
                                    text_column = col
                                    break
                            
                            if text_column:
                                for idx, row in df.iterrows():
                                    resume_text = str(row[text_column])
                                    if resume_text.strip():
                                        candidate_name = str(row.get('name', f'Candidate_{idx+1}'))
                                        candidate = self.scoring_engine.analyze_resume(
                                            resume_text, candidate_name, st.session_state.job_analysis
                                        )
                                        candidates.append(candidate)
                            else:
                                st.error("❌ CSV file must contain a 'resume_text', 'text', 'content', 'resume', or 'description' column")
                                return
                        
                        elif dataset_file.type == "application/json":
                            # Handle JSON file with resume data
                            json_data = json.loads(dataset_file.read().decode('utf-8'))
                            
                            if isinstance(json_data, list):
                                for idx, item in enumerate(json_data):
                                    if isinstance(item, dict):
                                        # Look for resume text in various fields
                                        resume_text = ""
                                        for field in ['resume_text', 'text', 'content', 'resume', 'description']:
                                            if field in item:
                                                resume_text = str(item[field])
                                                break
                                        
                                        if resume_text.strip():
                                            candidate_name = str(item.get('name', f'Candidate_{idx+1}'))
                                            candidate = self.scoring_engine.analyze_resume(
                                                resume_text, candidate_name, st.session_state.job_analysis
                                            )
                                            candidates.append(candidate)
                            else:
                                st.error("❌ JSON file must contain an array of resume objects")
                                return
                        
                        # Rank candidates
                        if candidates:
                            ranked_candidates = self.scoring_engine.rank_candidates(
                                candidates, st.session_state.job_analysis
                            )
                            st.session_state.candidates.extend(ranked_candidates)
                            
                            st.success(f"✅ Successfully processed {len(candidates)} resumes from dataset!")
                            st.rerun()
                        else:
                            st.warning("⚠️ No valid resumes found in dataset")
                            
                    except Exception as e:
                        st.error(f"Error processing dataset: {e}")
        
        else:
            # Fallback text input mode
            st.info("📝 Paste resume texts below (separate with '---')")
            resume_texts = st.text_area(
                "Paste resume texts (one per resume, separated by '---'):",
                height=300,
                placeholder="Paste resume content here. Separate multiple resumes with '---' on a new line."
            )
            
            if st.button("🚀 Analyze Resumes", type="primary"):
                if resume_texts.strip():
                    with st.spinner("Processing resumes..."):
                        try:
                            # Split resumes by separator
                            resume_list = [text.strip() for text in resume_texts.split('---') if text.strip()]
                            
                            candidates = []
                            for i, resume_text in enumerate(resume_list):
                                # Analyze resume
                                candidate = self.scoring_engine.analyze_resume(
                                    resume_text,
                                    f"Resume_{i+1}",
                                    st.session_state.job_analysis
                                )
                                candidates.append(candidate)
                            
                            # Rank candidates
                            if candidates:
                                ranked_candidates = self.scoring_engine.rank_candidates(
                                    candidates, st.session_state.job_analysis
                                )
                                st.session_state.candidates = ranked_candidates
                                
                                st.success(f"✅ Successfully analyzed {len(ranked_candidates)} resumes!")
                                
                                # Show top candidates preview
                                st.subheader("🏆 Top Candidates")
                                for candidate in ranked_candidates[:3]:
                                    score_class = self.get_score_class(candidate['final_score'])
                                    candidate_html = f"""
                                    <div class="candidate-card">
                                        <h3>#{candidate['rank']} {candidate['name']}</h3>
                                        <p><strong>Score:</strong> <span class="score-{score_class}">{candidate['final_score']:.1f}/100</span></p>
                                        <p><strong>Email:</strong> {candidate['email']}</p>
                                        <p><strong>Skills Match:</strong> {candidate['skill_coverage']['overall_coverage']:.1f}%</p>
                                        <p><strong>Experience:</strong> {candidate['years_of_experience']} years</p>
                                    </div>
                                    """
                                    st.markdown(candidate_html, unsafe_allow_html=True)
                            else:
                                st.error("No resumes could be processed successfully")
                        
                        except Exception as e:
                            st.error(f"Error processing resumes: {e}")
            else:
                st.error("Please enter resume texts")
    
    def get_score_class(self, score):
        """Get CSS class for score styling."""
        if score >= 70:
            return "high"
        elif score >= 50:
            return "medium"
        else:
            return "low"
    
    def render_ranking_dashboard(self):
        """Render the ranking dashboard."""
        st.header("📊 Ranking Dashboard")
        
        if not st.session_state.candidates:
            st.warning("⚠️ No candidates analyzed yet. Please upload resumes first.")
            return
        
        # Summary metrics with red bar structure
        total_candidates = len(st.session_state.candidates)
        avg_score = sum(c['final_score'] for c in st.session_state.candidates) / len(st.session_state.candidates)
        high_scored = len([c for c in st.session_state.candidates if c['final_score'] >= 70])
        top_score = st.session_state.candidates[0]['final_score']
        
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #dc3545 0%, #c62828 100%); 
                    padding: 0.5rem; border-radius: 8px; margin: 0.3rem 0; 
                    box-shadow: 0 1px 4px rgba(220,53,69,0.3);">
            <div style="display: flex; justify-content: space-around; align-items: center;">
                <div style="text-align: center; color: white;">
                    <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                        👥
                    </div>
                    <div style="font-size: 1.1rem; font-weight: bold;">
                        {total_candidates}
                    </div>
                    <div style="font-size: 0.7rem; opacity: 0.9;">
                        Total Candidates
                    </div>
                </div>
                <div style="text-align: center; color: white;">
                    <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                        📈
                    </div>
                    <div style="font-size: 1.1rem; font-weight: bold;">
                        {avg_score:.1f}
                    </div>
                    <div style="font-size: 0.7rem; opacity: 0.9;">
                        Average Score
                    </div>
                </div>
                <div style="text-align: center; color: white;">
                    <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                        🎯
                    </div>
                    <div style="font-size: 1.1rem; font-weight: bold;">
                        {high_scored}
                    </div>
                    <div style="font-size: 0.7rem; opacity: 0.9;">
                        High Scoring (70+)
                    </div>
                </div>
                <div style="text-align: center; color: white;">
                    <div style="font-size: 0.9rem; font-weight: bold; margin-bottom: 0.2rem;">
                        🏆
                    </div>
                    <div style="font-size: 1.1rem; font-weight: bold;">
                        {top_score:.1f}
                    </div>
                    <div style="font-size: 0.7rem; opacity: 0.9;">
                        Top Score
                    </div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Detailed rankings table
        st.subheader("📋 Detailed Rankings")
        
        # Prepare data for table
        table_data = []
        for candidate in st.session_state.candidates:
            table_data.append({
                'Rank': candidate.get('rank', len(table_data) + 1),
                'Name': candidate['name'],
                'Score': f"{candidate['final_score']:.1f}",
                'Semantic Match': f"{candidate['semantic_similarity']:.1f}%",
                'Skill Coverage': f"{candidate['skill_coverage']['overall_coverage']:.1f}%",
                'Experience': f"{candidate['years_of_experience']} years",
                'Email': candidate['email']
            })
        
        df = pd.DataFrame(table_data)
        st.dataframe(df, use_container_width=True)
        
        # Ranking chart
        st.markdown("""
        <style>
        div[data-testid="stVerticalBlock"] > div:has(div[data-testid="stSubheader"]) + div[data-testid="stVerticalBlock"] {
            margin-top: -2.5rem !important;
        }
        </style>
        """, unsafe_allow_html=True)
        st.subheader("📈 Candidate Rankings")
        ranking_chart = self.visualizer.create_ranking_bar_chart(st.session_state.candidates)
        st.plotly_chart(ranking_chart, use_container_width=True)
        
        # Download results
        st.subheader("💾 Export Results")
        
        col1, col2, col3 = st.columns([1, 3, 1])
        with col1:
            # Generate CSV data once
            csv_data = self.export_to_csv()
            st.download_button(
                label="📥 Download CSV",
                data=csv_data,
                file_name="ranking_results.csv",
                mime="text/csv"
            )
        
        with col2:
            # Empty column for spacing
            pass
        
        with col3:
            # Generate detailed PDF report
            pdf_data = self.export_detailed_report()
            if pdf_data:
                st.download_button(
                    label="📄 Download Detailed Report",
                    data=pdf_data,
                    file_name="detailed_analysis_report.pdf",
                    mime="application/pdf"
                )
            else:
                st.warning("PDF generation unavailable")
    
    def render_candidate_insights(self):
        """Render candidate insights page."""
        st.header("📈 Candidate Insights")
        
        if not st.session_state.candidates:
            st.warning("⚠️ No candidates analyzed yet. Please upload resumes first.")
            return
        
        # Candidate selector
        candidate_names = [f"#{c['rank']} {c['name']}" for c in st.session_state.candidates]
        selected_candidate_idx = st.selectbox(
            "Select a candidate to analyze:",
            range(len(candidate_names)),
            format_func=lambda x: candidate_names[x]
        )
        
        selected_candidate = st.session_state.candidates[selected_candidate_idx]
        
        # Candidate overview
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""
            <div style='background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%); 
                        padding: 15px; border-radius: 8px; 
                        color: white; text-align: center; 
                        box-shadow: 0 4px 12px rgba(220, 53, 69, 0.3);'>
                <div style='font-size: 0.9rem; margin-bottom: 5px;'>Overall Score</div>
                <div style='font-size: 1.8rem; font-weight: bold;'>{:.1f}/100</div>
            </div>
            """.format(selected_candidate['final_score']), unsafe_allow_html=True)
        with col2:
            st.markdown("""
            <div style='background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%); 
                        padding: 15px; border-radius: 8px; 
                        color: white; text-align: center; 
                        box-shadow: 0 4px 12px rgba(220, 53, 69, 0.3);'>
                <div style='font-size: 0.9rem; margin-bottom: 5px;'>Skill Coverage</div>
                <div style='font-size: 1.8rem; font-weight: bold;'>{:.1f}%</div>
            </div>
            """.format(selected_candidate['skill_coverage']['overall_coverage']), unsafe_allow_html=True)
        with col3:
            st.markdown("""
            <div style='background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%); 
                        padding: 15px; border-radius: 8px; 
                        color: white; text-align: center; 
                        box-shadow: 0 4px 12px rgba(220, 53, 69, 0.3);'>
                <div style='font-size: 0.9rem; margin-bottom: 5px;'>Experience</div>
                <div style='font-size: 1.8rem; font-weight: bold;'>{:.1f} years</div>
            </div>
            """.format(selected_candidate['years_of_experience']), unsafe_allow_html=True)
        
        # Score breakdown
        st.subheader("📊 Score Breakdown")
        
        # Create tabular format for scores
        exp_years = selected_candidate['years_of_experience']
        min_req_years = selected_candidate.get('experience_requirements', {}).get('minimum_years', 1)
        exp_score = min(exp_years / min_req_years * 100, 100)
        
        score_data = [
            {"Component": "Skill Match", "Score": f"{selected_candidate['skill_coverage']['overall_coverage']:.1f}%", "Weight": "60%"},
            {"Component": "Experience", "Score": f"{exp_score:.1f}%", "Weight": "25%"},
            {"Component": "Keyword Match", "Score": f"{selected_candidate['semantic_similarity']:.1f}%", "Weight": "15%"}
        ]
        score_df = pd.DataFrame(score_data)
        st.dataframe(score_df, use_container_width=True, hide_index=True)
        
        # Skill analysis
        # Skills are stored in skill_coverage, not as a separate 'skills' key
        skill_coverage = selected_candidate['skill_coverage']
        if skill_coverage.get('matched_skills'):
            st.subheader("✅ Matched Skills")
            st.write(", ".join(skill_coverage['matched_skills']))
        
        if skill_coverage.get('missing_skills'):
            st.subheader("❌ Missing Skills")
            st.write(", ".join(skill_coverage['missing_skills']))
        
        # Show skill coverage by category if available
        if skill_coverage.get('by_category'):
            st.subheader("📊 Skill Coverage by Category")
            for category, data in skill_coverage['by_category'].items():
                st.write(f"**{category}**: {data['total_matched']}/{data['total_required']} matched ({data['score']:.1f}%)")
        
        # Detailed explanation
        if selected_candidate.get('explanation'):
            st.subheader("🔍 Detailed Explanation")
            st.markdown(selected_candidate['explanation'])
    
    def render_bias_analysis(self):
        """Render bias analysis page."""
        st.header("⚖️ Bias Analysis")
        
        # Job description bias
        if st.session_state.job_analysis:
            st.subheader("📋 Job Description Bias Analysis")
            job_bias = st.session_state.job_analysis['bias_analysis']
            
            # Enhanced bias details with gradient styling
            st.markdown("""
            <style>
                .bias-metric {
                    background: linear-gradient(135deg, #dc3545 0%, #c62828 50%, #8b0000 100%);
                    padding: 20px;
                    border-radius: 10px;
                    color: white;
                    text-align: center;
                    box-shadow: 0 6px 20px rgba(220, 53, 69, 0.3);
                    margin-bottom: 15px;
                }
                .bias-label {
                    font-size: 0.9rem;
                    margin-bottom: 8px;
                    opacity: 0.9;
                }
                .bias-value {
                    font-size: 1.8rem;
                    font-weight: bold;
                }
                .bias-issues {
                    background-color: #000000;
                    padding: 20px;
                    border-radius: 8px;
                    border-left: 4px solid #dc3545;
                    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
                    margin-top: 15px;
                }
                .bias-recommendations {
                    background-color: #000000;
                    padding: 20px;
                    border-radius: 8px;
                    border-left: 4px solid #28a745;
                    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.3);
                    margin-top: 15px;
                }
            </style>
            """, unsafe_allow_html=True)
            
            # Bias metrics in enhanced format (side by side)
            col1, col2 = st.columns(2)
            with col1:
                bias_key = 'overall_risk_level' if 'overall_risk_level' in job_bias else 'risk_level'
                score_key = 'overall_risk_score' if 'overall_risk_score' in job_bias else 'risk_score'
                
                # Risk Level with gradient and conditional coloring
                risk_level = job_bias[bias_key]
                if risk_level.lower() == 'low':
                    risk_color = "#28a745"  # Green for low risk
                    risk_text_color = "white"
                elif risk_level.lower() == 'medium':
                    risk_color = "#ffc107"  # Yellow for medium risk
                    risk_text_color = "black"
                else:  # high risk
                    risk_color = "#dc3545"  # Red for high risk
                    risk_text_color = "white"
                
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, {risk_color} 0%, {risk_color}dd 50%, {risk_color}99 100%);
                            padding: 20px; border-radius: 10px; 
                            color: {risk_text_color}; text-align: center; 
                            box-shadow: 0 6px 20px rgba({risk_color.replace('#', '')}, 0.3);
                            margin-bottom: 15px;'>
                    <div class='bias-label'>Risk Level</div>
                    <div class='bias-value'>{risk_level}</div>
                </div>
                """, unsafe_allow_html=True)
                
                # Risk Score with gradient and conditional coloring
                risk_score = job_bias[score_key]
                if risk_score <= 40:
                    score_color = "#28a745"  # Green for 0-40
                    score_text_color = "white"
                elif risk_score <= 70:
                    score_color = "#ffc107"  # Yellow for 50-70
                    score_text_color = "black"
                else:  # 80-100
                    score_color = "#dc3545"  # Red for 80-100
                    score_text_color = "white"
                
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, {score_color} 0%, {score_color}dd 50%, {score_color}99 100%);
                            padding: 20px; border-radius: 10px; 
                            color: {score_text_color}; text-align: center; 
                            box-shadow: 0 6px 20px rgba({score_color.replace('#', '')}, 0.3);
                            margin-bottom: 15px;'>
                    <div class='bias-label'>Risk Score</div>
                    <div class='bias-value'>{risk_score:.1f}/100</div>
                </div>
                """, unsafe_allow_html=True)
            
            with col2:
                # Detected Issues (top section)
                if job_bias.get('detected_items'):
                    all_issues = " • " + " • ".join(job_bias['detected_items'])
                    st.markdown(f"""
                    <div class='bias-issues'>
                        <strong style='color: #dc3545; font-size: 1.1rem;'>⚠️ Detected Issues:</strong>
                        <div style='margin-top: 10px; color: #dc3545; font-size: 1rem; line-height: 1.5;'>
                            {all_issues}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                
                # Recommendations (bottom section)
                if job_bias.get('recommendations'):
                    all_recommendations = " • " + " • ".join(job_bias['recommendations'])
                    st.markdown(f"""
                    <div class='bias-recommendations' style='margin-top: 20px;'>
                        <strong style='color: #28a745; font-size: 1.2rem; font-weight: bold;'>💡 Recommendations</strong>
                        <div style='margin-top: 10px; color: #28a745; font-size: 1rem; line-height: 1.5;'>
                            {all_recommendations}
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
        
        st.divider()
        
        # Resume bias analysis
        if st.session_state.candidates:
            st.subheader("📄 Resume Bias Analysis")
            
            # Overall bias distribution
            bias_scores = []
            for c in st.session_state.candidates:
                bias = c['bias_analysis']
                score_key = 'overall_risk_score' if 'overall_risk_score' in bias else 'risk_score'
                bias_scores.append(bias[score_key])
            
            # Enhanced bias distribution table
            bias_data = []
            for i, c in enumerate(st.session_state.candidates):
                bias = c['bias_analysis']
                score_key = 'overall_risk_score' if 'overall_risk_score' in bias else 'risk_score'
                level_key = 'overall_risk_level' if 'overall_risk_level' in bias else 'risk_level'
                
                # Determine risk level color
                risk_score = bias[score_key]
                if risk_score <= 30:
                    risk_color = "🟢 Low"
                elif risk_score <= 60:
                    risk_color = "🟡 Medium"
                else:
                    risk_color = "🔴 High"
                
                bias_data.append({
                    "Rank": f"#{c['rank']}",
                    "Candidate Name": c['name'],
                    "Risk Score": f"{risk_score:.1f}/100",
                    "Risk Level": risk_color
                })
            
            bias_df = pd.DataFrame(bias_data)
            
            # Custom styling for the table
            st.markdown("""
            <style>
                .enhanced-table {
                    border-radius: 8px;
                    overflow: hidden;
                    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
                }
                .dataframe {
                    font-size: 14px;
                }
            </style>
            """, unsafe_allow_html=True)
            
            st.dataframe(bias_df, use_container_width=True, hide_index=True, 
                        column_config={
                            "Rank": st.column_config.TextColumn("Rank", width="small"),
                            "Candidate Name": st.column_config.TextColumn("Candidate Name", width="large"),
                            "Risk Score": st.column_config.TextColumn("Risk Score", width="medium"),
                            "Risk Level": st.column_config.TextColumn("Risk Level", width="medium")
                        })
    
    def render_resume_comparison(self):
        """Render resume comparison page."""
        st.header("🔍 Resume Comparison")
        
        if len(st.session_state.candidates) < 2:
            st.warning("⚠️ Need at least 2 candidates to compare. Please upload more resumes.")
            return
        
        # Candidate selectors
        candidate_names = [f"#{c['rank']} {c['name']}" for c in st.session_state.candidates]
        
        col1, col2 = st.columns(2)
        with col1:
            candidate1_idx = st.selectbox(
                "Select first candidate:",
                range(len(candidate_names)),
                format_func=lambda x: candidate_names[x],
                key="candidate1"
            )
        
        with col2:
            candidate2_idx = st.selectbox(
                "Select second candidate:",
                range(len(candidate_names)),
                format_func=lambda x: candidate_names[x],
                key="candidate2"
            )
        
        if candidate1_idx == candidate2_idx:
            st.error("Please select different candidates for comparison")
            return
        
        candidate1 = st.session_state.candidates[candidate1_idx]
        candidate2 = st.session_state.candidates[candidate2_idx]
        
        # Comparison chart
        st.subheader("📊 Side-by-Side Comparison")
        comparison_chart = self.visualizer.create_candidate_comparison_chart(candidate1, candidate2)
        st.plotly_chart(comparison_chart, use_container_width=True)
        
        # Winner announcement
        winner = candidate1['name'] if candidate1['final_score'] > candidate2['final_score'] else candidate2['name']
        score_diff = abs(candidate1['final_score'] - candidate2['final_score'])
        st.success(f"🏆 **{winner}** ranks higher with a score difference of {score_diff:.1f} points")
    
    def get_score_class(self, score):
        """Get CSS class for score styling."""
        if score >= 70:
            return "high"
        elif score >= 50:
            return "medium"
        else:
            return "low"
    
    def export_to_csv(self):
        """Export candidate rankings to CSV."""
        data = []
        for i, candidate in enumerate(st.session_state.candidates, 1):
            data.append({
                'Rank': candidate.get('rank', i),
                'Name': candidate['name'],
                'Email': candidate['email'],
                'Phone': candidate['phone'],
                'Final_Score': candidate['final_score'],
                'Semantic_Similarity': candidate['semantic_similarity'],
                'Skill_Coverage': candidate['skill_coverage']['overall_coverage'],
                'Years_Experience': candidate['years_of_experience'],
                'Bias_Risk_Score': candidate['bias_analysis']['overall_risk_score'],
                'Bias_Risk_Level': candidate['bias_analysis']['overall_risk_level']
            })
        
        df = pd.DataFrame(data)
        return df.to_csv(index=False)
    
    def export_detailed_report(self):
        """Export detailed analysis report to PDF format."""
        if not PDF_GENERATION_AVAILABLE:
            st.error("PDF generation not available. Please install reportlab: pip install reportlab")
            return None
            
        from io import BytesIO
        
        # Create PDF buffer
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading_style = styles['Heading1']
        subheading_style = styles['Heading2']
        normal_style = styles['Normal']
        
        # Create custom styles
        custom_style = ParagraphStyle(
            'Custom',
            parent=normal_style,
            fontSize=10,
            spaceAfter=6,
            leftIndent=20
        )
        
        # Title
        story.append(Paragraph("AI Resume Ranking System - Detailed Analysis Report", title_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph(f"<b>Generated on:</b> {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        story.append(Spacer(1, 20))
        
        # Job Description Analysis Section
        story.append(Paragraph("JOB DESCRIPTION ANALYSIS", heading_style))
        story.append(Spacer(1, 12))
        
        if st.session_state.job_analysis:
            job_analysis = st.session_state.job_analysis
            
            # Job Analysis Summary Table
            job_data = [
                ['Metric', 'Value'],
                ['Bias Risk Score', f"{job_analysis['bias_analysis']['overall_risk_score']:.1f}/100"],
                ['Bias Risk Level', job_analysis['bias_analysis']['overall_risk_level']],
                ['Skill Categories', str(len(job_analysis['skills']))]
            ]
            
            job_table = Table(job_data, colWidths=[2*inch, 3*inch])
            job_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(job_table)
            story.append(Spacer(1, 20))
            
            # Required Skills Section
            if job_analysis['skills']:
                story.append(Paragraph("Required Skills by Category", subheading_style))
                story.append(Spacer(1, 12))
                
                for category, skills in job_analysis['skills'].items():
                    if skills:
                        story.append(Paragraph(f"<b>{category}:</b> {', '.join(skills)}", custom_style))
                story.append(Spacer(1, 20))
        
        # Candidate Analysis Section
        story.append(Paragraph("CANDIDATE ANALYSIS", heading_style))
        story.append(Spacer(1, 12))
        
        for candidate in st.session_state.candidates:
            # Candidate Header
            story.append(Paragraph(f"Rank #{candidate['rank']}: {candidate['name']}", subheading_style))
            story.append(Spacer(1, 8))
            
            # Candidate Details Table
            candidate_data = [
                ['Attribute', 'Value'],
                ['Email', candidate['email']],
                ['Phone', candidate['phone']],
                ['Final Score', f"{candidate['final_score']:.1f}/100"],
                ['Semantic Similarity', f"{candidate['semantic_similarity']:.1f}%"],
                ['Skill Coverage', f"{candidate['skill_coverage']['overall_coverage']:.1f}%"],
                ['Years of Experience', str(candidate['years_of_experience'])],
                ['Bias Risk Score', f"{candidate['bias_analysis']['overall_risk_score']:.1f}/100"],
                ['Bias Risk Level', candidate['bias_analysis']['overall_risk_level']]
            ]
            
            candidate_table = Table(candidate_data, colWidths=[2*inch, 3*inch])
            candidate_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(candidate_table)
            story.append(Spacer(1, 12))
            
            # Skills Breakdown
            if candidate['skill_coverage']:
                story.append(Paragraph("Skills by Category", normal_style))
                story.append(Spacer(1, 6))
                
                for category, coverage in candidate['skill_coverage'].items():
                    if category != 'overall_coverage':
                        if isinstance(coverage, dict):
                            match_rate = coverage.get('match_rate', 0)
                            matched_skills = coverage.get('matched_skills', [])
                            if matched_skills:
                                story.append(Paragraph(f"<b>{category}:</b> {match_rate:.1f}% - {', '.join(matched_skills)}", custom_style))
                            else:
                                story.append(Paragraph(f"<b>{category}:</b> {match_rate:.1f}%", custom_style))
                        else:
                            story.append(Paragraph(f"<b>{category}:</b> {coverage}", custom_style))
                story.append(Spacer(1, 15))
        
        # Summary Section
        story.append(Paragraph("SUMMARY", heading_style))
        story.append(Spacer(1, 12))
        
        if st.session_state.candidates:
            avg_score = sum(c['final_score'] for c in st.session_state.candidates) / len(st.session_state.candidates)
            high_score = max(c['final_score'] for c in st.session_state.candidates)
            qualified = len([c for c in st.session_state.candidates if c['final_score'] >= 70])
            
            summary_data = [
                ['Metric', 'Value'],
                ['Total Candidates', str(len(st.session_state.candidates))],
                ['Average Score', f"{avg_score:.1f}/100"],
                ['Highest Score', f"{high_score:.1f}/100"],
                ['Qualified Candidates (70+)', str(qualified)]
            ]
            
            summary_table = Table(summary_data, colWidths=[2*inch, 3*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(summary_table)
            story.append(Spacer(1, 20))
        
        # Detailed Rankings Table Section
        story.append(Paragraph("DETAILED RANKINGS TABLE", heading_style))
        story.append(Spacer(1, 12))
        
        if st.session_state.candidates:
            # Create rankings table data
            rankings_data = [['Rank', 'Name', 'Email', 'Final Score', 'Skill Coverage', 'Experience', 'Bias Risk Level']]
            
            for candidate in sorted(st.session_state.candidates, key=lambda x: x['final_score'], reverse=True):
                rankings_data.append([
                    str(candidate['rank']),
                    candidate['name'],
                    candidate['email'],
                    f"{candidate['final_score']:.1f}/100",
                    f"{candidate['skill_coverage']['overall_coverage']:.1f}%",
                    f"{candidate['years_of_experience']} years",
                    candidate['bias_analysis']['overall_risk_level']
                ])
            
            # Create rankings table
            rankings_table = Table(rankings_data, colWidths=[0.8*inch, 2*inch, 2*inch, 1.2*inch, 1.2*inch, 1.2*inch, 1.2*inch])
            rankings_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
            ]))
            story.append(rankings_table)
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()


# Main execution
if __name__ == "__main__":
    app = ResumeRankingApp()
    app.run()
